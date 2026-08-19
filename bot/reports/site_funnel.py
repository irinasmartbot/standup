"""TG site funnel report: afisha_besplat + quick_booking cohort."""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone, timedelta
from typing import Any, Optional

import psycopg
from psycopg.rows import dict_row

from bot.config import BOOKINGS_SOURCE, DATABASE_URL

logger = logging.getLogger(__name__)

MSK = timezone(timedelta(hours=3))

_PERSON_KEY = """
COALESCE(
    ae.user_id::text,
    CASE WHEN ae.telegram_id IS NOT NULL THEN 'tg:' || ae.telegram_id::text END
)
""".strip()

_FUNNEL_STEPS = [
    ("site_entry", "1. Перешли с сайта (afisha_besplat + quick_booking)", None),
    ("branch_proverka", "2. Зашли в «Проверку» (бесплатная бронь)", {"name": "branch_proverka"}),
    (
        "browse_proverka",
        "3. Смотрели афишу (карточки проверки)",
        {"any_of": ["show_card"], "format": "proverka"},
    ),
    (
        "booking_created",
        "4. Создали бронь (проверка)",
        {"name": "booking_created", "format": "proverka"},
    ),
    (
        "booking_confirmed",
        "5. Получили билет (подтвердили бронь)",
        {"name": "booking_confirmed", "format": "proverka"},
    ),
    (
        "booking_cancelled",
        "6. Отменили бронь",
        {"name": "booking_cancelled", "format": "proverka"},
    ),
    (
        "booking_annulled",
        "7. Бронь аннулирована",
        {"name": "booking_annulled", "format": "proverka"},
    ),
]


def _use_postgres() -> bool:
    return BOOKINGS_SOURCE == "postgres" and bool(DATABASE_URL)


def _event_predicate(step_filter: dict | None) -> str:
    if step_filter is None:
        return "FALSE"
    if "name" in step_filter:
        clause = "ae.name = %(ev_name)s"
        if step_filter.get("format"):
            clause += " AND ae.props->>'format' = %(ev_format)s"
        return clause
    if "any_of" in step_filter:
        names = step_filter["any_of"]
        placeholders = ", ".join(f"'{n}'" for n in names)
        clause = f"ae.name IN ({placeholders})"
        if step_filter.get("format"):
            clause += " AND ae.props->>'format' = %(ev_format)s"
        return clause
    return "FALSE"


def _fetch_report(conn) -> dict[str, Any]:
    person_in_cohort = _PERSON_KEY.replace("ae.", "s.")
    with conn.cursor(row_factory=dict_row) as cur:
        cur.execute(
            """
            SELECT
                MIN(created_at) AS min_at,
                MAX(created_at) AS max_at
            FROM analytics_events
            """
        )
        bounds = dict(cur.fetchone() or {})

        cur.execute(
            f"""
            CREATE TEMP TABLE site_cohort ON COMMIT DROP AS
            SELECT
                {person_in_cohort} AS person_key,
                MIN(s.user_id) AS user_id,
                MIN(s.telegram_id) AS telegram_id,
                MIN(s.created_at) AS first_site_at,
                BOOL_OR(s.props->>'payload' = 'afisha_besplat') AS via_afisha,
                BOOL_OR(s.props->>'payload' = 'quick_booking') AS via_quick
            FROM analytics_events s
            WHERE s.name = 'bot_start'
              AND s.channel = 'telegram'
              AND s.props->>'payload' IN ('afisha_besplat', 'quick_booking')
              AND s.telegram_id IS NOT NULL
            GROUP BY {person_in_cohort}
            """
        )

        cur.execute("SELECT COUNT(*)::int AS n FROM site_cohort")
        site_total = int(cur.fetchone()["n"])
        cur.execute("SELECT COUNT(*)::int AS n FROM site_cohort WHERE via_afisha")
        via_afisha = int(cur.fetchone()["n"])
        cur.execute("SELECT COUNT(*)::int AS n FROM site_cohort WHERE via_quick")
        via_quick = int(cur.fetchone()["n"])

        steps: list[dict] = []
        prev_uniques = site_total
        for key, label, flt in _FUNNEL_STEPS:
            if key == "site_entry":
                uniques = site_total
            else:
                params: dict[str, Any] = {}
                if flt and flt.get("name"):
                    params["ev_name"] = flt["name"]
                if flt and flt.get("format"):
                    params["ev_format"] = flt["format"]
                pred = _event_predicate(flt)
                cur.execute(
                    f"""
                    SELECT COUNT(DISTINCT sc.person_key)::int AS uniques
                    FROM site_cohort sc
                    WHERE EXISTS (
                        SELECT 1
                        FROM analytics_events ae
                        WHERE {_PERSON_KEY} = sc.person_key
                          AND ae.channel = 'telegram'
                          AND ae.created_at >= sc.first_site_at
                          AND {pred}
                    )
                    """,
                    params,
                )
                uniques = int(cur.fetchone()["uniques"] or 0)

            conv_prev = (
                round(100.0 * uniques / prev_uniques, 1)
                if prev_uniques and key != "site_entry"
                else None
            )
            conv_site = round(100.0 * uniques / site_total, 1) if site_total else None
            steps.append(
                {
                    "key": key,
                    "label": label,
                    "uniques": uniques,
                    "conv_prev": conv_prev,
                    "conv_site": conv_site,
                }
            )
            if key not in ("booking_cancelled", "booking_annulled") and uniques:
                prev_uniques = uniques

        cur.execute(
            """
            SELECT
                COUNT(DISTINCT b.id)::int AS bookings,
                COUNT(DISTINCT b.user_id)::int AS users
            FROM bookings b
            INNER JOIN site_cohort sc ON sc.user_id = b.user_id
            WHERE b.format = 'proverka'
              AND b.source = 'telegram'
            """
        )
        db_created = dict(cur.fetchone() or {})
        cur.execute(
            """
            SELECT COUNT(DISTINCT b.user_id)::int AS users
            FROM bookings b
            INNER JOIN site_cohort sc ON sc.user_id = b.user_id
            WHERE b.format = 'proverka'
              AND b.source = 'telegram'
              AND b.status = 'confirmed'
            """
        )
        db_confirmed_users = int(cur.fetchone()["users"] or 0)
        cur.execute(
            """
            SELECT COUNT(DISTINCT b.user_id)::int AS users
            FROM bookings b
            INNER JOIN site_cohort sc ON sc.user_id = b.user_id
            WHERE b.format = 'proverka'
              AND b.source = 'telegram'
              AND b.status = 'cancelled'
            """
        )
        db_cancelled_users = int(cur.fetchone()["users"] or 0)
        cur.execute(
            """
            SELECT COUNT(DISTINCT b.user_id)::int AS users
            FROM bookings b
            INNER JOIN site_cohort sc ON sc.user_id = b.user_id
            WHERE b.format = 'proverka'
              AND b.source = 'telegram'
              AND b.status = 'annulled'
            """
        )
        db_annulled_users = int(cur.fetchone()["users"] or 0)

    min_at = bounds.get("min_at")
    max_at = bounds.get("max_at")
    period_label = "весь период с начала учёта analytics"
    if min_at and max_at:
        d0 = min_at.astimezone(MSK).strftime("%d.%m.%Y")
        d1 = max_at.astimezone(MSK).strftime("%d.%m.%Y")
        period_label = f"{d0} — {d1} (МСК)"

    return {
        "available": True,
        "generated_at": datetime.now(MSK).strftime("%d.%m.%Y %H:%M"),
        "period_label": period_label,
        "site_total": site_total,
        "via_afisha": via_afisha,
        "via_quick": via_quick,
        "steps": steps,
        "db_crosscheck": {
            "bookings": db_created.get("bookings") or 0,
            "users_created": db_created.get("users") or 0,
            "users_confirmed": db_confirmed_users,
            "users_cancelled": db_cancelled_users,
            "users_annulled": db_annulled_users,
        },
    }


def fetch_site_funnel_report() -> dict[str, Any]:
    empty: dict[str, Any] = {"available": False}
    if not _use_postgres():
        return empty
    try:
        with psycopg.connect(DATABASE_URL) as conn:
            return _fetch_report(conn)
    except Exception:
        logger.exception("fetch_site_funnel_report failed")
        return empty


def _fmt_pct(value: Optional[float]) -> str:
    if value is None:
        return "—"
    return f"{value:.1f}%"


def build_site_funnel_docx_bytes(report: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("Воронка Telegram: переходы с сайта", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Сформировано: {report['generated_at']} (МСК)")
    doc.add_paragraph(f"Период данных: {report['period_label']}")
    doc.add_paragraph(
        "Когорта: пользователи Telegram, у которых был /start с payload "
        "afisha_besplat или quick_booking (ссылки с сайта). "
        "Дальнейшие шаги считаются после первого такого захода."
    )
    doc.add_paragraph(
        "Канал: только Telegram. VK и обычный /start без payload в эту воронку не входят."
    )

    doc.add_heading("Сводка по источникам захода", level=1)
    t0 = doc.add_table(rows=4, cols=2)
    t0.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Уникальных с сайта (объединённо)", str(report["site_total"])),
            ("из них заход afisha_besplat (хотя бы раз)", str(report["via_afisha"])),
            ("из них заход quick_booking (хотя бы раз)", str(report["via_quick"])),
            (
                "Примечание",
                "Один человек мог зайти по обеим ссылкам — в объединённой строке он один раз.",
            ),
        ]
    ):
        t0.rows[i].cells[0].text = a
        t0.rows[i].cells[1].text = b

    doc.add_heading("Воронка «Проверка материала»", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Этап"
    hdr[1].text = "Уникальных людей"
    hdr[2].text = "Конверсия от предыдущего шага"
    hdr[3].text = "Конверсия от захода с сайта"

    for step in report["steps"]:
        row = table.add_row().cells
        row[0].text = step["label"]
        row[1].text = str(step["uniques"])
        row[2].text = _fmt_pct(step["conv_prev"])
        row[3].text = _fmt_pct(step["conv_site"])

    doc.add_heading("Проверка по таблице броней (когорта с сайта)", level=1)
    doc.add_paragraph(
        "Дублирующая сверка по bookings: те же user_id, format=proverka, source=telegram."
    )
    xc = report["db_crosscheck"]
    t1 = doc.add_table(rows=5, cols=2)
    t1.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Броней создано (записей)", str(xc["bookings"])),
            ("Уникальных с хотя бы одной бронью", str(xc["users_created"])),
            ("Уникальных с билетом (status=confirmed)", str(xc["users_confirmed"])),
            ("Уникальных с отменой (status=cancelled)", str(xc["users_cancelled"])),
            ("Уникальных с аннуляцией (status=annulled)", str(xc["users_annulled"])),
        ]
    ):
        t1.rows[i].cells[0].text = a
        t1.rows[i].cells[1].text = b

    doc.add_heading("Методология", level=1)
    doc.add_paragraph(
        "• События берутся из analytics_events (PostgreSQL).\n"
        "• «Перешли с сайта» — bot_start с payload afisha_besplat или quick_booking.\n"
        "• «Получили билет» — событие booking_confirmed (нажали «Получить билет» / подтвердили).\n"
        "• Отмена и аннуляция — отдельные люди; один пользователь может иметь несколько броней.\n"
        "• Импорт из Salebot / ручная заливка без bot_start с сайта в когорту не попадает."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
