#!/usr/bin/env python3
"""Разовый отчёт TG-воронки с сайта. Запуск на сервере из /home/standup/app:

    venv/bin/pip install python-docx -q
    venv/bin/python scripts/site_funnel_standalone.py

Файл: data/reports/site_funnel_tg.docx
"""

from __future__ import annotations

import io
import os
import subprocess
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

try:
    from dotenv import load_dotenv

    load_dotenv(ROOT / ".env")
except ImportError:
    pass

MSK = timezone(timedelta(hours=3))
OUT = ROOT / "data" / "reports" / "site_funnel_tg.docx"

PERSON_KEY = """
COALESCE(
    ae.user_id::text,
    CASE WHEN ae.telegram_id IS NOT NULL THEN 'tg:' || ae.telegram_id::text END
)
""".strip()

STEPS = [
    ("site_entry", "1. Перешли с сайта (afisha_besplat + quick_booking)", None),
    ("branch_proverka", "2. Зашли в «Проверку» (бесплатная бронь)", {"name": "branch_proverka"}),
    (
        "browse_proverka",
        "3. Смотрели афишу (карточки проверки)",
        {"any_of": ["show_card"], "format": "proverka"},
    ),
    ("booking_created", "4. Создали бронь (проверка)", {"name": "booking_created", "format": "proverka"}),
    (
        "booking_confirmed",
        "5. Получили билет (подтвердили бронь)",
        {"name": "booking_confirmed", "format": "proverka"},
    ),
    ("booking_cancelled", "6. Отменили бронь", {"name": "booking_cancelled", "format": "proverka"}),
    ("booking_annulled", "7. Бронь аннулирована", {"name": "booking_annulled", "format": "proverka"}),
]


def _pred(flt):
    if not flt:
        return "FALSE"
    if "name" in flt:
        s = "ae.name = %(ev_name)s"
        if flt.get("format"):
            s += " AND ae.props->>'format' = %(ev_format)s"
        return s
    if "any_of" in flt:
        ph = ", ".join(f"'{n}'" for n in flt["any_of"])
        s = f"ae.name IN ({ph})"
        if flt.get("format"):
            s += " AND ae.props->>'format' = %(ev_format)s"
        return s
    return "FALSE"


def _ensure_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        return Document, WD_ALIGN_PARAGRAPH
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        return Document, WD_ALIGN_PARAGRAPH


def fetch(conn):
    import psycopg
    from psycopg.rows import dict_row

    pk_cohort = PERSON_KEY.replace("ae.", "s.")
    with conn.cursor(row_factory=dict_row) as cur:
        cur.execute("SELECT MIN(created_at) min_at, MAX(created_at) max_at FROM analytics_events")
        bounds = dict(cur.fetchone() or {})
        cur.execute(
            f"""
            CREATE TEMP TABLE site_cohort ON COMMIT DROP AS
            SELECT {pk_cohort} person_key, MIN(s.user_id) user_id,
                   MIN(s.created_at) first_site_at,
                   BOOL_OR(s.props->>'payload' = 'afisha_besplat') via_afisha,
                   BOOL_OR(s.props->>'payload' = 'quick_booking') via_quick
            FROM analytics_events s
            WHERE s.name = 'bot_start' AND s.channel = 'telegram'
              AND s.props->>'payload' IN ('afisha_besplat', 'quick_booking')
              AND s.telegram_id IS NOT NULL
            GROUP BY {pk_cohort}
            """
        )
        cur.execute("SELECT COUNT(*)::int n FROM site_cohort")
        total = int(cur.fetchone()["n"])
        cur.execute("SELECT COUNT(*)::int n FROM site_cohort WHERE via_afisha")
        via_afisha = int(cur.fetchone()["n"])
        cur.execute("SELECT COUNT(*)::int n FROM site_cohort WHERE via_quick")
        via_quick = int(cur.fetchone()["n"])

        rows = []
        prev = total
        for key, label, flt in STEPS:
            if key == "site_entry":
                u = total
            else:
                p = {}
                if flt and flt.get("name"):
                    p["ev_name"] = flt["name"]
                if flt and flt.get("format"):
                    p["ev_format"] = flt["format"]
                cur.execute(
                    f"""
                    SELECT COUNT(DISTINCT sc.person_key)::int uniques
                    FROM site_cohort sc
                    WHERE EXISTS (
                      SELECT 1 FROM analytics_events ae
                      WHERE {PERSON_KEY} = sc.person_key
                        AND ae.channel = 'telegram'
                        AND ae.created_at >= sc.first_site_at
                        AND {_pred(flt)}
                    )
                    """,
                    p,
                )
                u = int(cur.fetchone()["uniques"] or 0)
            cp = round(100 * u / prev, 1) if prev and key != "site_entry" else None
            cs = round(100 * u / total, 1) if total else None
            rows.append({"label": label, "uniques": u, "conv_prev": cp, "conv_site": cs})
            if key not in ("booking_cancelled", "booking_annulled") and u:
                prev = u

        cur.execute(
            """
            SELECT COUNT(DISTINCT b.id)::int bookings, COUNT(DISTINCT b.user_id)::int users
            FROM bookings b JOIN site_cohort sc ON sc.user_id = b.user_id
            WHERE b.format = 'proverka' AND b.source = 'telegram'
            """
        )
        xc0 = dict(cur.fetchone() or {})
        for st in ("confirmed", "cancelled", "annulled"):
            cur.execute(
                f"""
                SELECT COUNT(DISTINCT b.user_id)::int users
                FROM bookings b JOIN site_cohort sc ON sc.user_id = b.user_id
                WHERE b.format = 'proverka' AND b.source = 'telegram' AND b.status = '{st}'
                """
            )
            xc0[f"users_{st}"] = int(cur.fetchone()["users"] or 0)

    mn, mx = bounds.get("min_at"), bounds.get("max_at")
    period = "весь период с начала учёта analytics"
    if mn and mx:
        period = f"{mn.astimezone(MSK).strftime('%d.%m.%Y')} — {mx.astimezone(MSK).strftime('%d.%m.%Y')} (МСК)"

    return {
        "generated_at": datetime.now(MSK).strftime("%d.%m.%Y %H:%M"),
        "period_label": period,
        "site_total": total,
        "via_afisha": via_afisha,
        "via_quick": via_quick,
        "steps": rows,
        "xc": xc0,
    }


def build_docx(r):
    Document, WD_ALIGN_PARAGRAPH = _ensure_docx()
    doc = Document()
    t = doc.add_heading("Воронка Telegram: переходы с сайта", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Сформировано: {r['generated_at']} (МСК)")
    doc.add_paragraph(f"Период данных: {r['period_label']}")
    doc.add_paragraph(
        "Когорта: Telegram /start с payload afisha_besplat или quick_booking. "
        "Шаги после первого такого захода. Только TG."
    )

    doc.add_heading("Сводка по источникам", 1)
    tb = doc.add_table(4, 2)
    tb.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Уникальных с сайта", str(r["site_total"])),
            ("afisha_besplat (хотя бы раз)", str(r["via_afisha"])),
            ("quick_booking (хотя бы раз)", str(r["via_quick"])),
            ("Примечание", "Один человек по обеим ссылкам — один раз в объединённой строке."),
        ]
    ):
        tb.rows[i].cells[0].text, tb.rows[i].cells[1].text = a, b

    doc.add_heading("Воронка «Проверка материала»", 1)
    tbl = doc.add_table(1, 4)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = (
        "Этап",
        "Уникальных",
        "Конверсия от пред. шага",
        "Конверсия от сайта",
    )
    for s in r["steps"]:
        row = tbl.add_row().cells
        row[0].text = s["label"]
        row[1].text = str(s["uniques"])
        row[2].text = "—" if s["conv_prev"] is None else f"{s['conv_prev']:.1f}%"
        row[3].text = "—" if s["conv_site"] is None else f"{s['conv_site']:.1f}%"

    doc.add_heading("Сверка по bookings", 1)
    x = r["xc"]
    t2 = doc.add_table(5, 2)
    t2.style = "Table Grid"
    for i, (a, b) in enumerate(
        [
            ("Броней (записей)", str(x.get("bookings", 0))),
            ("Уникальных с бронью", str(x.get("users", 0))),
            ("С билетом (confirmed)", str(x.get("users_confirmed", 0))),
            ("С отменой (cancelled)", str(x.get("users_cancelled", 0))),
            ("С аннуляцией (annulled)", str(x.get("users_annulled", 0))),
        ]
    ):
        t2.rows[i].cells[0].text, t2.rows[i].cells[1].text = a, b

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    url = os.getenv("DATABASE_URL")
    if not url:
        print("ERROR: DATABASE_URL not set (.env in /home/standup/app)", file=sys.stderr)
        return 1
    import psycopg

    with psycopg.connect(url) as conn:
        report = fetch(conn)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_bytes(build_docx(report))
    print(f"OK: {OUT}")
    print(f"Site cohort: {report['site_total']}")
    for s in report["steps"]:
        print(f"  {s['label']}: {s['uniques']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
